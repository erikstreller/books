# custom_invitations.py

from pathlib import Path

import docx


def invite(path: Path) -> None:
    text_file = "guest_list.txt"
    word_file = "invitation_cards.docx"

    text_path = path / text_file
    word_path = path / word_file

    doc = docx.Document(str(word_path))
    guest_list = open(str(text_path))

    for guest in guest_list:
        doc.add_paragraph(
            "It would be a pleasure to have the company of", style="style_serif"
        )
        doc.add_paragraph(guest.replace("\n", " "), style="style_name")
        doc.add_paragraph("at 11010 Memory Lane on the Evening of", style="style_serif")
        doc.add_paragraph("April 1st", style="style_date")
        doc.add_paragraph("at 7 o'clock", style="style_serif")
        doc.add_page_break()

    doc.save(str(word_path))


if __name__ == "__main__":
    path = Path(__file__).parent.resolve()

    invite(path)
